import streamlit as st
from google.genai.errors import APIError
from utils import get_gemini_client
import io, time
import json
from docx import Document
import usage_manager as um

if "performance_log" not in st.session_state:
    st.session_state.performance_log = []
if 'quiz_data' not in st.session_state:
    st.session_state.quiz_data = None
if 'quiz_submitted' not in st.session_state:
    st.session_state.quiz_submitted = False
if 'quiz_score' not in st.session_state:
    st.session_state.quiz_score = 0

model_name = "gemini-2.5-flash"
client = get_gemini_client()

st.title("❓ Smart Quiz Generator")
st.markdown("Generate interactive quizzes with instant grading and explanations.")

with st.form("quiz_generator_form"):
    quiz_topic = st.text_input("Topic to Quiz on", placeholder="e.g., Newton's Laws of Motion")
    num_questions = st.selectbox("Number of Questions", options=[5, 10, 15], index=0)
    quiz_type = st.selectbox("Question Type", options=["Multiple Choice", "True/False"], index=0)
    difficulty = st.slider("Difficulty Level", 1, 5, 3, help="1=Easy, 5=Hard")

    generate_quiz_clicked = st.form_submit_button("Generate Quiz", type="primary")


if generate_quiz_clicked:
    # Reset previous quiz state
    st.session_state.quiz_data = None
    st.session_state.quiz_submitted = False
    st.session_state.quiz_score = 0

    if not quiz_topic:
        st.error("Please enter a topic.")
        st.stop()

    if not um.check_guest_limit("Quiz Generator", limit=1):
        st.page_link("pages/00_login.py", label="Login/Signup", icon="🔑")
        st.stop()

    DIFFICULTY_MAP = {1: "introductory", 2: "beginner", 3: "intermediate", 4: "advanced", 5: "expert"}

    QUIZ_PROMPT = f"""
    You are an expert quiz creator. Create a {quiz_type} quiz on the topic: "{quiz_topic}".
    Difficulty: {DIFFICULTY_MAP[difficulty]}.
    Number of Questions: {num_questions}.

    OUTPUT FORMAT:
    Return ONLY a raw JSON list of dictionaries. Do NOT use Markdown code blocks (like ```json).
    Each dictionary must have these keys:
    - "question": The question text
    - "options": A list of strings (e.g., ["Option A", "Option B", "Option C", "Option D"] or ["True", "False"])
    - "answer": The exact string of the correct option
    - "explanation": A short explanation of why it is correct
    """

    with st.spinner(f"Generating a {num_questions} question quiz on {quiz_topic}"):
        try:
            response = client.models.generate_content(
                model=model_name,
                contents=[QUIZ_PROMPT]
            )

            cleaned_text = response.text.replace("```json", "").replace("```", "").strip()
            st.session_state.quiz_data = json.loads(cleaned_text)
            st.rerun()

        except json.JSONDecodeError:
            st.error("Error: AI output invalid JSON. Please try again.")
        except APIError as e:
            error_text = str(e)
            if "429" in error_text or "RESOURCE_EXHAUSTED" in error_text.upper():
                if "api_key" in st.session_state:
                    del st.session_state.api_key
                st.error("🚨 **Quota Exceeded!** The Gemini API key has hit its limit")
                st.stop()
            elif "503" in error_text:
                st.markdown("The Gemini AI model is currently experiencing high traffic. Please try again later.")
                st.info("Meanwhile, try other non-AI features like GPA Calculator, Study Scheduler, etc.")
            else:
                st.error(f"An API error occurred during generation: {e}")
        except Exception as e:
            st.error(f"Error: {e}")


if st.session_state.quiz_data:
    if st.button("Generate New Quiz"):
        st.session_state.quiz_data = None
        st.session_state.quiz_submitted = False
        st.session_state.quiz_score = 0
        st.rerun()


if st.session_state.quiz_data:
    st.divider()
    st.subheader(f"📝 Quiz: {quiz_topic if quiz_topic else 'Generated Quiz'}")

    with st.form("taking_quiz_form"):
        user_answers = {}
        for idx, q in enumerate(st.session_state.quiz_data):
            st.markdown(f"**{idx + 1}. {q['question']}**")
            user_answers[idx] = st.radio(
                "Select Answer:",
                q['options'],
                index=None,
                key=f"q_{idx}",
                label_visibility="collapsed"
            )
            st.markdown("---")

        submit_quiz = st.form_submit_button("Submit & Grade")

    if submit_quiz:
        st.session_state.quiz_submitted = True
        score = 0
        for idx, q in enumerate(st.session_state.quiz_data):
            if user_answers[idx] == q['answer']:
                score += 1
        st.session_state.quiz_score = score

    # --- Log performance ---
        if "user" in st.session_state:
            user_id = st.session_state.user.id
            um.log_performance(
                user_id=user_id,
                feature="Quiz Generator",
                score=score,
                total_questions=num_questions,
                correct_answers=score
            )


    if st.session_state.quiz_submitted:
        total = len(st.session_state.quiz_data)
        pct = (st.session_state.quiz_score / total) * 100

        st.session_state.performance_log.append({
            "type": "quiz",
            "topic": quiz_topic,
            "total": num_questions,
            "score": st.session_state.quiz_score,
            "percentage": pct,
            "difficulty": difficulty,
            "questions": num_questions,
            "timestamp": time.time()
        })

        auth_user_id = st.session_state.user.id
        um.log_usage(auth_user_id, "Quiz Generator", "generated", {"topic": quiz_topic})

        if pct >= 80:
            st.balloons()
            st.success(f"### Excellent! You scored {st.session_state.quiz_score}/{total} ({pct:.0f}%)")
        elif pct >= 50:
            st.warning(f"### Good job! You scored {st.session_state.quiz_score}/{total} ({pct:.0f}%)")
        else:
            st.error(f"### Keep practicing. You scored {st.session_state.quiz_score}/{total} ({pct:.0f}%)")

        st.subheader("🔍 Answer Key & Explanations")
        for idx, q in enumerate(st.session_state.quiz_data):
            user_choice = st.session_state.get(f"q_{idx}")
            with st.expander(f"Q{idx + 1}: {q['question']}", expanded=True):
                if user_choice == q['answer']:
                    st.success(f"✅ Your Answer: {user_choice}")
                else:
                    st.error(f"❌ Your Answer: {user_choice}")
                    st.info(f"✅ Correct Answer: {q['answer']}")
                st.markdown(f"**Explanation:** {q['explanation']}")


        doc = Document()
        doc.add_heading(f"Quiz Results: {quiz_topic}", 0)
        doc.add_paragraph(f"Final Score: {st.session_state.quiz_score}/{total}\n")
        for idx, q in enumerate(st.session_state.quiz_data):
            doc.add_heading(f"Q{idx + 1}: {q['question']}", level=2)
            doc.add_paragraph(f"Options: {q['options']}")
            doc.add_paragraph(f"Correct Answer: {q['answer']}")
            doc.add_paragraph(f"Explanation: {q['explanation']}")
            doc.add_paragraph("-" * 20)

        doc_io = io.BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)


        if um.premium_gate("Download Quiz Results"):
            download_clicked = st.download_button(
                label="Download Results as DOCX",
                data=doc_io,
                file_name=f"{quiz_topic} Quiz_Results.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
            if download_clicked:
                # Delete quiz data immediately after download
                st.session_state.quiz_data = None
                st.session_state.quiz_submitted = False
                st.session_state.quiz_score = 0
        else:
            st.info("Creating an account is free and saves your progress!")
            st.page_link("pages/00_login.py", label="Login/Signup", icon="🔑")
