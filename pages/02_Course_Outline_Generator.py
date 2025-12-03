import streamlit as st
from google.genai.errors import APIError
from utils import get_gemini_client
from docx import Document
import io
import usage_manager as um

model_name = "gemini-2.5-flash"

st.title("📝 Course Outline Generator")
st.markdown("Instantly generate a detailed, university-level course syllabus and outline")

client = get_gemini_client()


if "course_outline" not in st.session_state:
    st.session_state.course_outline = None
if "course_outline_filename" not in st.session_state:
    st.session_state.course_outline_filename = None


if st.session_state.course_outline:

    st.success("📌 A previously generated course outline was found.")

    st.markdown("### ▶ Download Previously Generated Outline")

    if um.premium_gate("Download Course Outline"):

        # Prepare DOCX again from stored text
        doc = Document()
        doc.add_heading(st.session_state.course_outline_filename.replace("_", " ").replace(".docx", ""), 0)
        doc.add_paragraph(st.session_state.course_outline)

        doc_io = io.BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)

        download_clicked = st.download_button(
            label="⬇ Download Course Outline (DOCX)",
            data=doc_io,
            file_name=st.session_state.course_outline_filename,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key="download_previous"
        )

        if download_clicked:
            del st.session_state.course_outline
            del st.session_state.course_outline_filename
            st.success("✔ Previous content cleared after download.")
            st.rerun()

    else:
        st.info("Create an account to download your previously generated outline.")
        st.page_link("pages/00_login.py", label="Login/Signup", icon="🔑")

    st.markdown("---")
    if st.button("🆕 Generate New Course Outline"):
        del st.session_state.course_outline
        del st.session_state.course_outline_filename
        st.rerun()

    st.stop()


with st.form("course_outline_form"):
    course_full_name = st.text_input("Course Full Name", placeholder="e.g. Introduction to Computer Science")
    course_code = st.text_input("Course Code (Optional)", placeholder="e.g., CSC 101")
    university_name = st.text_input("University Name (Optional)",
                                    placeholder="e.g., Harvard University or Lagos State University")

    submitted = st.form_submit_button("Generate Outline", type="primary")

if submitted:
    if not course_full_name:
        st.error("Please enter the **Course Full Name** to generate an outline.")
        st.stop()

    if not um.check_guest_limit("Course Outline Generator", limit=1):
        st.page_link("pages/00_login.py", label="Login/Signup", icon="🔑")
        st.stop()

    # Construct prompt
    uni_context = f"taught at {university_name}." if university_name else "taught at a major Nigerian University."
    code_context = f"(Code: {course_code})" if course_code else ""

    OUTLINE_PROMPT = f"""
    You are an expert curriculum designer. Generate a comprehensive, 12-week university_level course outline for the 
    course: "{course_full_name}" {code_context}. The course should reflect standards {uni_context}.

    **REQUIRED SECTIONS:**
    1. **Course Description:** (2-3 sentences)
    2. **Course Objectives:** (4-5 bullet points)
    3. **12-Week Schedule:** Use a **Markdown table** with the columns: **Week**, **Topic**, and **Key Learning Objectives**.

    Ensure the output is formatted cleanly using **Markdown**.
    """

    with st.spinner(f"Generating 12-week outline for {course_full_name}..."):
        try:
            response = client.models.generate_content(
                model=model_name,
                contents=[OUTLINE_PROMPT]
            )

            outline_text = response.text

            # Save to session
            st.session_state.course_outline = outline_text
            filename = f"{course_full_name.replace(' ', '_')}_Outline.docx"
            st.session_state.course_outline_filename = filename

            # Display
            st.subheader("✔ Generated Course Outline")
            st.markdown(outline_text)

            # Prepare DOCX for download
            doc = Document()
            doc.add_heading(f"Course Outline: {course_full_name}", 0)
            doc.add_paragraph(outline_text)
            doc_io = io.BytesIO()
            doc.save(doc_io)
            doc_io.seek(0)

            if um.premium_gate("Download Course Outline"):
                download_clicked = st.download_button(
                    label="⬇ Download as DOCX",
                    data=doc_io,
                    file_name=filename,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key="download_generated"
                )

                if download_clicked:
                    del st.session_state.course_outline
                    del st.session_state.course_outline_filename
                    st.success("✔ Content cleared after download.")
                    st.rerun()
            else:
                st.info("Creating an account is free and saves your progress!")
                st.page_link("pages/00_login.py", label="Login/Signup", icon="🔑")

        except APIError as e:
            msg = str(e)
            if "429" in msg or "RESOURCE_EXHAUSTED" in msg.upper():
                st.error("🚨 **Quota Exceeded!** The Gemini API key has hit its limit.")
                if "api_key" in st.session_state:
                    del st.session_state.api_key
                st.stop()
            elif "503" in msg:
                st.warning("The Gemini AI model is currently experiencing high traffic. Please try again later.")
            else:
                st.error(f"API Error: {e}")

        except Exception as e:
            st.error(f"Unexpected Error: {e}")

st.markdown("---")
if st.button("🆕 Generate New Course Outline"):
    if "course_outline" in st.session_state:
        del st.session_state.course_outline
    if "course_outline_filename" in st.session_state:
        del st.session_state.course_outline_filename
    st.rerun()
