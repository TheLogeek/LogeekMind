import streamlit as st
from google.genai.errors import APIError
from utils import get_gemini_client
import io
from docx import Document

model_name = "gemini-2.5-flash"

st.title("❓ Smart Quiz Generator")
st.markdown("Instantly generate a custom quiz on any topic, complete with answers and explanations.")

client = get_gemini_client()

# Quiz Input Form
with st.form("quiz_generator_form"):
    quiz_topic = st.text_input(
        "Topic to Quiz on",
        placeholder="e.g., Newton's Laws of Motion"
    )
    num_questions = st.selectbox(
        "Number of Questions",
        options=[5, 10, 15],
        index=0
    )
    quiz_type = st.selectbox(
        "Question Type",
        options=["Multple Choice", "True/False", "Short Answer"],
        index=0
    )
    difficulty = st.slider(
        "Difficulty Level",
        min_value=1,
        max_value=5,
        value=3,
        help="1 = Easy (Introductory), 5 = Hard (Advanced Concepts)"
    )

    submitted = st.form_submit_button("Generate Quiz", type="primary")

if submitted:
    if not quiz_topic:
        st.error("Please enter a **Topic to Quiz On**.")
        st.stop()

    # AI prompt

    DIFFICULTY_MAP = {
        1: "very simple/introductory level",
        2: "beginner level",
        3: "intermediate level",
        4: "advanced level",
        5: "expert/challenging level"
    }

    QUIZ_PROMPT = f"""
    You are a professional quiz master. Your task is to generate a comprehensive quiz.
    
    **QUIZ SPECIFICATIONS:**
    * **Topic:** {quiz_topic}
    * **Number of Questions:** {num_questions}
    * **Question Type:** {quiz_type}
    * **Difficulty:** {DIFFICULTY_MAP[difficulty]}
    
    **REQUIRED FORMAT:**
    1. Generate all the questions sequentially.
    2. Immediately after the last question, start a new section titled **'---ANSWER KEY & EXPLANATIONS---'**.
    3. In the answer key section, list the correct answer for each question (e.g., '1. A' or '1. True' or '1. [Short 
    Answer]')
    
    Ensure the entire output is formatted cleanly using **Markdown** for lists and bolding.
    """

    with st.spinner(f"Generating a {num_questions} question quiz on {quiz_topic}..."):
        try:
            response = client.models.generate_content(
                model=model_name,
                contents=[QUIZ_PROMPT]
            )
            quiz_text = response.text

            st.subheader(f"🧠 Your {quiz_type} Quiz: {quiz_topic}")

            quiz_sections = quiz_text.split('--- ANSWER KEY & EXPLANATIONS ---')
            if len(quiz_sections) > 1:
                st.markdown(quiz_sections[0]) #questions

                with st.expander("Show Answers & Explanations"):
                    st.markdown("### Answer Key & Explanations")
                    st.markdown(quiz_sections[1].strip())
            else:
                st.markdown(quiz_text)

            doc = Document()
            doc.add_heading(f"{quiz_type}Quiz: {quiz_topic}", 0)
            doc.add_paragraph(quiz_text)

            doc_io = io.BytesIO()
            doc.save(doc_io)
            doc_io.seek(0)

            st.download_button(
                label="Download as DOCX",
                data=doc_io,
                file_name=f"{quiz_topic} Quiz.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

        except APIError as e:
            error_text = str(e)
            if "429" in error_text or "RESOURCE_EXHAUSTED" in error_text.upper():
                if "api_key" in st.session_state:
                    del st.session_state.api_key
                st.error("🚨 **Quota Exceeded!** The shared key has hit its limit. Please enter your own Gemini API "
                         "Key in the sidebar to continue.")
                st.stop()
            elif "503" in error_text:
                st.markdown("The Gemini AI model is currently experiencing high traffic. Please try again later. "
                            "Thank you for your patience!")
                st.info("In the meantime, you can try other non-AI features **(GPA Calculator, Study Scheduler Lecture Note-to-Audio Converter, Lecture Audio-to-Text Converter)**")
            else:
                st.error(f"An API error occurred during generation: {e}")

        except Exception as e:
            st.error(f"An unexpected error occurred: {e}")