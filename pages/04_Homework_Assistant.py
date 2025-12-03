import streamlit as st
from PIL import Image
from google.genai.errors import APIError
from utils import get_gemini_client
from docx import Document
import io, time
import usage_manager as um

model_name = "gemini-2.5-flash"

st.title("📸 Homework Assistant")
st.markdown("Upload a picture of your homework problem and get a step-by-step, downloadable solution.")

client = get_gemini_client()


if "hw_image" not in st.session_state:
    st.session_state.hw_image = None
if "hw_context" not in st.session_state:
    st.session_state.hw_context = ""
if "hw_solution" not in st.session_state:
    st.session_state.hw_solution = None
if "hw_doc" not in st.session_state:
    st.session_state.hw_doc = None


uploaded_file = st.file_uploader(
    "Upload Image of Homework Problem",
    type=["jpg", "jpeg", "png"]
)

if uploaded_file is not None:
    st.session_state.hw_image = Image.open(uploaded_file)

if st.session_state.hw_image is not None:
    st.image(st.session_state.hw_image, caption="Uploaded homework problem", use_column_width=True)


st.session_state.hw_context = st.text_area(
    "Add Context (Optional)",
    value=st.session_state.hw_context,
    placeholder="e.g., This is a kinematics problem or I'm stuck on Step 3.",
    help="Give the AI any extra information to help it solve the problem."
)

def generate_solution():
    if not um.check_guest_limit("Homework Assistant", limit=1):
        st.page_link("pages/00_login.py", label="Login/Signup", icon="🔑")
        st.stop()

    try:
        full_prompt = f"""
        You are a rigorous academic solver. Based on the image and the user's instructions (if any),
        provide a complete and accurate solution. The output must be in a clean, professional, and easily 
        readable **Markdown format**.

        Instructions:
        {st.session_state.hw_context if st.session_state.hw_context else "The user provided no instructions"}
        """

        with st.spinner("Generating Solution..."):
            response = client.models.generate_content(
                model=model_name,
                contents=[st.session_state.hw_image, full_prompt]
            )

        st.session_state.hw_solution = response.text

        # Build DOCX
        doc = Document()
        doc.add_heading("Homework Solution", 0)
        doc.add_heading(f"Instructions: {st.session_state.hw_context}", 1)
        doc.add_paragraph(response.text)

        doc_io = io.BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)
        st.session_state.hw_doc = doc_io

    except APIError as e:
        error_text = str(e)
        if "429" in error_text or "RESOURCE_EXHAUSTED" in error_text.upper():
            if "api_key" in st.session_state:
                del st.session_state.api_key
            st.error("🚨 **Quota Exceeded!** The Gemini API key has hit its limit")
            st.stop()
        elif "503" in error_text:
            st.warning("Gemini is experiencing heavy traffic. Try again shortly.")
            st.info("Meanwhile, explore other non-AI tools like GPA Calculator or Study Scheduler.")
        else:
            st.error(f"API Error: {e}")
    except Exception as e:
        st.error(f"Unexpected Error: {e}")

col1, col2 = st.columns(2)

with col1:
    if st.button("Generate Solution", type="primary"):
        if st.session_state.hw_image is None:
            st.warning("Please upload an image first.")
        else:
            generate_solution()

with col2:
    if st.session_state.hw_solution:
        if st.button("Generate New"):
            st.session_state.hw_solution = None
            st.session_state.hw_doc = None

if st.session_state.hw_solution:
    st.subheader("Generated Solution (Preview)")
    st.markdown(st.session_state.hw_solution)

    if um.premium_gate("Download Homework Solution"):
        st.download_button(
            label="Download Solution as DOCX",
            data=st.session_state.hw_doc,
            file_name=f"homework_solution_{time.strftime('%Y%m%d%H%M')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            on_click=lambda: (
                st.session_state.update({"hw_solution": None, "hw_doc": None})
            )
        )
    else:
        st.info("Create a free account to download and save your generated solutions.")
        st.page_link("pages/00_login.py", label="Login/Signup", icon="🔑")
