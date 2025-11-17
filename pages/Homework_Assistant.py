import streamlit as st
from PIL import Image
from google.genai.errors import APIError
from utils import get_gemini_client
from docx import Document
import io, time

model_name = "gemini-2.5-flash"

st.title("📸 Homework Assistant")
st.markdown("Upload a picture of your homework problem and get a step-by-step, downloadable solution.")

client = get_gemini_client()

uploaded_file = st.file_uploader("Upload Image of Homework Problem", type=["jpg", "jpeg", "png"])

if uploaded_file is not None:
    problem_image = Image.open(uploaded_file)

    st.image(problem_image, caption="Uploaded homework problem", use_column_width=True)

    user_prompt = st.text_area(
        "Add Context (Optional)",
        placeholder="e.g., This is a kinematics problem or I'm stuck on Step 3.",
        help="Give the AI any extra information to help it solve the problem."
    )

    if st.button("Generate Solution", type="primary"):
        try:

            full_prompt = (f"""You are a rigorous academic solver. Based on the image and the user's instructions(if 
            any),provide a complete and accurate solution. The output must be in a clean, professional, and easily 
            readable **Markdown format**. Instructions: 
            {user_prompt if user_prompt else "The user provided no instructions"}""")

            with st.spinner("Generating Solution...."):
                response = client.models.generate_content(
                model=model_name,
                contents=[problem_image, full_prompt]
            )

                st.subheader("Generated Solution (Preview)")
                solution_text = response.text
                st.markdown(solution_text)

                doc = Document()
                doc.add_heading('Homework Solution', 0)
                doc.add_heading(f"Instructions: {user_prompt}", 1)
                doc.add_paragraph(solution_text)

                doc_io = io.BytesIO()
                doc.save(doc_io)
                doc_io.seek(0)

                st.download_button(
                    label="Download Solution as DOCX",
                    data=doc_io,
                    file_name=f"homework_solution_{time.strftime('%Y%m%d%H%M')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

        except APIError as e:
            error_text = str(e)
            if "429" in error_text or "RESOURCE_EXHAUSTED" in error_text.upper():
                if "api_key" in st.session_state:
                    del st.session_state.api_key
                st.error(
                    "🚨 **Quota Exceeded!** The shared key has hit its limit. Please enter your own Gemini API "
                    "Key in the sidebar to continue.")
                st.stop()
            elif "503" in error_text:
                st.markdown("The Gemini AI model is currently experiencing high traffic. Please try again later. "
                            "Thank you for your patience!")
                st.info("In the meantime, you can try other non-AI features **(GPA Calculator, Study Scheduler Lecture Note-to-Audio Converter, Lecture Audio-to-Text Converter)**")
            else:
                st.error(f"An API error occurred during generation: {e}")

        except Exception as e:
            st.error(f"An unexpected error occurred: {e}")