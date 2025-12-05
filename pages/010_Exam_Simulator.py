import streamlit as st
import json
import io, time
from docx import Document
from streamlit_autorefresh import st_autorefresh
from google.genai.errors import APIError
from utils import get_gemini_client
import usage_manager as um

if "performance_log" not in st.session_state:
    st.session_state.performance_log = []
if 'exam_stage' not in st.session_state:
    st.session_state.exam_stage = "setup"  # Options: setup, active, finished
if 'exam_data' not in st.session_state:
    st.session_state.exam_data = []
if 'exam_answers' not in st.session_state:
    st.session_state.exam_answers = {}
if 'start_time' not in st.session_state:
    st.session_state.start_time = None
if 'duration_mins' not in st.session_state:
    st.session_state.duration_mins = 30
if 'exam_score' not in st.session_state:
    st.session_state.exam_score = 0

st.set_page_config(page_title="Exam Simulator", layout="wide")
st.title("🔥 Exam Simulator")

def _save_radio_answer(idx):
    key = f"q_{idx}"
    st.session_state["exam_answers"][str(idx)] = st.session_state.get(key)

def calculate_grade(score, total):
    percentage = (score / total) * 100
    if percentage >= 70:
        return "A", "Excellent! Distinction level."
    elif percentage >= 60:
        return "B", "Very Good. Keep it up."
    elif percentage >= 50:
        return "C", "Credit. You passed, but barely."
    elif percentage >= 45:
        return "D", "Pass. You need to study more."
    elif percentage >= 40:
        return "E", "Weak Pass. Dangerous territory."
    else:
        return "F", "Fail. You are not ready for this exam."

def clear_exam_session_state():
    for key in ["exam_stage", "exam_data", "exam_answers", "start_time", "duration_mins", "exam_score", "course_code"]:
        if key in st.session_state:
            del st.session_state[key]


model_name = "gemini-2.5-flash"
client = get_gemini_client()

# STAGE 1: SETUP
if st.session_state.exam_stage == "setup":
    st.markdown("### 📝 Setup Your Mock Exam")
    col1, col2 = st.columns(2)

    with col1:
        course_code = st.text_input("Course Name", placeholder="e.g., Introduction to Computer Science")
        topic = st.text_input("Specific Topic (Optional)", placeholder="e.g., Algorithms")

    with col2:
        duration = st.selectbox("Exam Duration", [1, 5, 10, 30, 60], index=2, format_func=lambda x: f"{x} Minutes")
        num_q = st.slider("Number of Questions", 5, 50, 20)

    if st.button("Start Exam ⏱️", type="primary"):
        if not um.check_guest_limit("Exam Simulator", limit=1):
            st.page_link("pages/00_login.py", label="Login/Signup", icon="🔑")
            st.stop()
        if not course_code:
            st.error("Please enter a Course Code.")
        else:
            with st.spinner("Prof. LogeekMind is preparing your exam papers..."):
                st.session_state.course_code = course_code
                st.session_state.exam_topic = topic
                prompt = f"""
                    You are a strict university professor setting a final exam.
                    Course: {course_code}
                    Topic: {topic}

                    Generate {num_q} HARD, examination-standard multiple-choice questions.
                    These should not be simple definitions. They should require critical thinking or application of concepts.

                    OUTPUT FORMAT:
                    Return ONLY a raw JSON list of dictionaries. Do NOT use Markdown code blocks.
                    Each dictionary must have these keys:
                    - "question": complex scenario or problem statement
                    - "options": A list of strings (e.g., ["Option A", "Option B", "Option C", "Option D"])
                    - "answer": The exact string of the correct option
                    - "explanation": A short explanation of why it is correct
                """
                try:
                    response = client.models.generate_content(model=model_name, contents=[prompt])
                    cleaned_text = response.text.replace("```json", "").replace("```", "").strip()
                    st.session_state.exam_data = json.loads(cleaned_text)
                except json.JSONDecodeError as e:
                    st.error(f"JSONDecodeError: AI output malformed: {e}")
                    st.stop()
                except APIError as e:
                    error_text = str(e)
                    if "429" in error_text or "RESOURCE_EXHAUSTED" in error_text.upper():
                        if "api_key" in st.session_state:
                            del st.session_state.api_key
                        st.error("🚨 Quota Exceeded! Gemini API limit reached.")
                    elif "503" in error_text:
                        st.error("Gemini AI model busy. Try again later.")
                    else:
                        st.error("API error during exam generation.")
                    st.stop()
                except Exception as e:
                    st.error(f"Unexpected error: {e}")
                    st.stop()

                # Set timers and stage
                st.session_state.duration_mins = duration
                st.session_state.start_time = time.time()
                st.session_state.exam_answers = {}
                st.session_state.exam_stage = "active"
                st.rerun()


# STAGE 2: ACTIVE EXAM
elif st.session_state.exam_stage == "active":
    elapsed_time = time.time() - st.session_state.start_time
    total_seconds = st.session_state.duration_mins * 60
    remaining_seconds = total_seconds - elapsed_time
    total_questions = len(st.session_state.exam_data)

    if total_questions == 0:
        st.warning("No exam data found. Please restart and generate the exam.")
    else:
        # Timer
        mins, secs = divmod(int(max(0, remaining_seconds)), 60)
        timer_color = "red" if mins < 2 else "green"
        st.markdown(f"""
            <div style="position: fixed; top: 60px; right: 20px; padding: 10px; background-color: {timer_color}; color: white; border-radius: 5px; z-index: 9999;">
                <b>Time Left: {mins:02d}:{secs:02d}</b>
            </div>
        """, unsafe_allow_html=True)

        # Progress
        answered_count = sum(1 for v in st.session_state["exam_answers"].values() if v is not None)
        progress_percent = answered_count / total_questions if total_questions > 0 else 0
        st.subheader("Question Progress")
        st.progress(progress_percent, text=f"**{answered_count} / {total_questions}** Questions Answered")
        st.write("---")

        # Questions
        for idx, q in enumerate(st.session_state.exam_data):
            st.markdown(f"**Q{idx + 1}: {q['question']}**")
            st.radio(
                "Select Answer:",
                q['options'],
                key=f"q_{idx}",
                label_visibility="collapsed",
                index=None,
                on_change=_save_radio_answer,
                args=(idx,)
            )
            st.write("---")

        # Manual submit
        if st.button("Submit Exam Now"):
            st.session_state.exam_stage = "finished"
            st.rerun()

        # Auto-submit when time ends
        if remaining_seconds <= 0:
            st.session_state.exam_stage = "finished"
            st.rerun()

        # Keep timer ticking
        st_autorefresh(interval=1000, key="timer_refresh")


# STAGE 3: RESULTS
elif st.session_state.exam_stage == "finished":
    st.title("📄 Exam Results")

    score = sum(
        1 for idx, q in enumerate(st.session_state.exam_data)
        if st.session_state["exam_answers"].get(str(idx)) == q.get('answer')
    )
    st.session_state.exam_score = score
    total = len(st.session_state.exam_data)
    grade, remark = calculate_grade(score, total)

    # --- Log performance ---
    if "user" in st.session_state:
        um.log_performance(
            user_id=st.session_state.user.id,
            feature="Exam Simulator",
            score=score,
            total_questions=total,
            correct_answers=score
        )

    st.markdown(f"""
        <div style="padding: 20px; background-color: #f0f2f6; border-radius: 10px; border-left: 10px solid {'#4CAF50' if grade in ['A','B'] else '#FF5722'};">
            <h2>Grade: {grade}</h2>
            <h3>Score: {score} / {total}</h3>
            <p><i>{remark}</i></p>
        </div><br>
    """, unsafe_allow_html=True)
    if grade in ["A", "B"]:
        st.balloons()

    st.session_state.performance_log.append({
        "type": "exam",
        "course": st.session_state.get("course_code", "course"),
        "topic": st.session_state.get("exam_topic", "topic"),
        "total": total,
        "score": score,
        "grade": grade,
        "duration": st.session_state.duration_mins,
        "timestamp": time.time()
    })

    if "user" in st.session_state:
        auth_user_id = st.session_state.user.id
        user_name = st.session_state.user_profile.username
        um.log_usage(auth_user_id, user_name, "Exam Simulator", "submitted_exam",
                 {"course": st.session_state.get("course_code", "course")})

    # Review answers
    with st.expander("View Corrections"):
        for idx, q in enumerate(st.session_state.exam_data):
            user_choice = st.session_state["exam_answers"].get(str(idx))
            with st.expander(f"Q{idx+1}: {q['question']}", expanded=True):
                if user_choice == q['answer']:
                    st.success(f"✅ Your Answer: {user_choice}")
                else:
                    st.error(f"❌ Your Answer: {user_choice or '(No answer)'}")
                    st.info(f"✅ Correct Answer: {q['answer']}")
                st.markdown(f"**Explanation:** {q.get('explanation', 'No explanation provided.')}")

    # Generate DOCX
    doc = Document()
    course_code = st.session_state.get("course_code", "course")
    doc.add_heading(f"Exam Results: {course_code}", 0)
    doc.add_paragraph(f"Final Score: {score}/{total}\nGrade: {grade}")

    for idx, q in enumerate(st.session_state.exam_data):
        doc.add_heading(f"Q{idx+1}: {q['question']}", level=2)
        doc.add_paragraph(f"Options: {q['options']}")
        doc.add_paragraph(f"Correct Answer: {q['answer']}")
        doc.add_paragraph(f"Explanation: {q['explanation']}")
        doc.add_paragraph("-"*20)

    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)


    if um.premium_gate("Download Exam Results"):
        st.download_button(
            label="Download Results as DOCX",
            data=doc_io,
            file_name=f"{course_code} Exam_Results.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            on_click=clear_exam_session_state
        )
    else:
        st.info("Creating an account is free and saves your progress!")
        st.page_link("pages/00_login.py", label="Login/Signup", icon="🔑")


    if st.button("Take Another Exam"):
        if um.premium_gate("Take Another Exam"):
            clear_exam_session_state()
            st.rerun()
        else:
            st.info("Creating an account is free and saves your progress!")
            st.page_link("pages/00_login.py", label="Login/Signup", icon="🔑")
